# ─────────────────────────────────────────────────────────────────────────────
#  doc_utils.py  –  python-docx utilities
# ─────────────────────────────────────────────────────────────────────────────

from __future__ import annotations
import copy, io
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ── Load / save ───────────────────────────────────────────────────────────────

def load_docx(file_like) -> Document:
    if isinstance(file_like, (io.BytesIO, io.RawIOBase, io.BufferedIOBase)):
        file_like.seek(0)
        return Document(file_like)
    try:
        file_like.seek(0)
    except Exception:
        pass
    return Document(io.BytesIO(file_like.read()))


def save_docx_to_bytes(doc: Document) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ── Text extraction ───────────────────────────────────────────────────────────

def extract_resume_text(doc: Document) -> str:
    """Plain-text dump of a resume .docx (body paragraphs + table cells)."""
    lines = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            lines.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    t = p.text.strip()
                    if t:
                        lines.append(t)
    return "\n".join(lines)


# ── Template context (sent to the AI) ─────────────────────────────────────────

def get_template_context(doc: Document) -> str:
    """
    Return a richly-annotated paragraph list for the AI prompt.

    Tab-separated lines (used for right-aligned two-column layout like
    "Company Name  →TAB→  Location") are explicitly marked so the AI
    knows to preserve the tab separator in its replacement values.

    Section headers and blank spacers are labelled so the AI knows not
    to replace them.
    """
    lines = []
    paragraphs = doc.paragraphs

    # Detect section headers: ALL-CAPS, no digits, 3+ chars
    def _is_header(text: str) -> bool:
        t = text.strip()
        return bool(t) and t.isupper() and len(t) >= 3 and not any(c.isdigit() for c in t)

    for i, para in enumerate(paragraphs):
        text = para.text

        if not text.strip():
            lines.append(f"[{i}]  ← BLANK SPACER — do not replace")
            continue

        if _is_header(text):
            lines.append(f"[{i}] {text}  ← SECTION HEADER — do not replace")
            continue

        if '\t' in text:
            left, _, right = text.partition('\t')
            lines.append(
                f"[{i}] {left.strip()}  →TAB→  {right.strip()}"
                f'  ← TAB-ALIGNED: return as "left text\\tright text"'
            )
            continue

        lines.append(f"[{i}] {text}")

    return "\n".join(lines)


# ── AI-generated paragraph replacements ──────────────────────────────────────

def apply_paragraph_replacements(doc: Document, replacements: dict) -> int:
    """
    Apply {paragraph_index: new_text} to *doc* in place.
    Works in descending order so insertions don't shift later indices.
    \\n in a value inserts multiple sibling paragraphs at that slot.
    Returns number of paragraphs touched.
    """
    paragraphs = doc.paragraphs
    applied = 0
    sorted_items = sorted(
        ((int(k), v) for k, v in replacements.items()
         if str(k).strip().lstrip("-").isdigit()),
        reverse=True,
    )
    for idx, new_text in sorted_items:
        if idx < 0 or idx >= len(paragraphs):
            continue
        para = paragraphs[idx]
        lines = new_text.split("\n")
        while lines and not lines[-1].strip():
            lines.pop()
        if not lines:
            _clear_para(para)
            applied += 1
            continue
        _set_para_text(para, lines[0])
        applied += 1
        anchor = para
        for extra in lines[1:]:
            new_elem = _clone_para(anchor, extra)
            anchor._element.addnext(new_elem)
    return applied


# ── User info injection (name, contact, company, date, signature) ─────────────

def apply_user_info(doc: Document, user_info: dict) -> None:
    """
    Replace template placeholder text with the user's real details.
    Uses text matching (not paragraph index) so it works with any template.

    user_info keys (all optional):
        name          – candidate's full name
        location      – e.g. "New York, NY"
        email         – e.g. "jane@email.com"
        phone         – e.g. "(555) 123-4567"
        linkedin      – e.g. "linkedin.com/in/janedoe"
        company       – target company name
        company_addr1 – company street address (optional)
        company_addr2 – company city/state/zip (optional)
        date          – letter date, e.g. "March 25, 2025"
    """
    name          = user_info.get("name", "").strip()
    location      = user_info.get("location", "").strip()
    email         = user_info.get("email", "").strip()
    phone         = user_info.get("phone", "").strip()
    linkedin      = user_info.get("linkedin", "").strip()
    github        = user_info.get("github", "").strip()
    company       = user_info.get("company", "").strip()
    company_addr1 = user_info.get("company_addr1", "").strip()
    company_addr2 = user_info.get("company_addr2", "").strip()
    date          = user_info.get("date", "").strip()

    # Build contact line from provided fields (GitHub appended after LinkedIn)
    contact_parts = [p for p in [location, email, phone, linkedin, github] if p]
    contact_line  = " | ".join(contact_parts) if contact_parts else ""

    # Known placeholder strings in the default templates
    _NAME_PLACEHOLDERS = {"Wolfie Seawolf"}
    _CONTACT_SIGNALS   = {
        "professional_email@gmail.com",
        "wolfie.seawolf@stonybrook.edu",
        "(123) 456",
        "(XXX) XXX",
    }
    _COMPANY_PLACEHOLDER  = "Company name"
    _ADDR1_PLACEHOLDER    = "XXXX Employer Rd."
    _ADDR2_PLACEHOLDER    = "New York, NY 11004"
    _DATE_PLACEHOLDER     = "January 1, 2024"
    _LOCATION_PLACEHOLDER = "Town, State"

    for para in doc.paragraphs:
        full = "".join(r.text for r in para.runs)
        stripped = full.strip()

        # ── Candidate name (standalone line) ─────────────────────────────────
        if name and stripped in _NAME_PLACEHOLDERS:
            _set_para_text(para, name)
            continue

        # ── Contact info line ─────────────────────────────────────────────────
        if contact_line and any(sig in full for sig in _CONTACT_SIGNALS):
            # The Stony Brook cover letter template has the date crammed onto
            # the same paragraph as the contact info. If that pattern is detected,
            # append the user's date to the contact line; otherwise just replace.
            if date and _DATE_PLACEHOLDER in full:
                _set_para_text(para, f"{contact_line}  {date}")
            else:
                _set_para_text(para, contact_line)
            continue

        # ── Location placeholder in resume header ─────────────────────────────
        if location and _LOCATION_PLACEHOLDER in stripped:
            _set_para_text(para, full.replace(_LOCATION_PLACEHOLDER, location))
            continue

        # ── Company name ──────────────────────────────────────────────────────
        if company and stripped == _COMPANY_PLACEHOLDER:
            _set_para_text(para, company)
            continue

        # ── Company address ───────────────────────────────────────────────────
        if company_addr1 and stripped == _ADDR1_PLACEHOLDER:
            _set_para_text(para, company_addr1)
            continue
        if company_addr2 and stripped == _ADDR2_PLACEHOLDER:
            _set_para_text(para, company_addr2)
            continue

        # ── Standalone date line ──────────────────────────────────────────────
        if date and stripped == _DATE_PLACEHOLDER:
            _set_para_text(para, date)
            continue


# ── Document preview (human-readable, not raw index pairs) ───────────────────

def build_document_preview(doc: Document) -> str:
    """
    Render the full document as clean plain text for display in the UI.
    ALL-CAPS section headers get a visual separator above them.
    """
    lines = []
    for para in doc.paragraphs:
        text    = para.text
        stripped = text.strip()

        if not stripped:
            lines.append("")
            continue

        # Section header detection: all-caps, at least 3 chars, no digits
        if (stripped.isupper()
                and len(stripped) >= 3
                and not any(c.isdigit() for c in stripped)):
            lines.append("")
            lines.append("─" * 48)
            lines.append(stripped)
            lines.append("─" * 48)
        else:
            lines.append(text)

    # Collapse runs of more than 2 blank lines
    result, prev_blank = [], 0
    for line in lines:
        if not line.strip():
            prev_blank += 1
            if prev_blank <= 1:
                result.append(line)
        else:
            prev_blank = 0
            result.append(line)

    return "\n".join(result).strip()


# ── Internal helpers ──────────────────────────────────────────────────────────

def _clear_para(para) -> None:
    for r in para.runs:
        r.text = ""

def _set_para_text(para, text: str) -> None:
    """
    Set the visible text of *para* to *text*, preserving run-level formatting.

    Special handling for tab-separated paragraphs (e.g. "Company\tLocation"):
    the tab character is preserved in the correct run so the Word tab-stop
    right-alignment continues to work.
    """
    if not para.runs:
        para.add_run(text)
        return

    if '\t' in text:
        left, _, right = text.partition('\t')

        # Find the run that currently owns the tab character
        tab_run_idx = next(
            (i for i, r in enumerate(para.runs) if '\t' in r.text),
            None,
        )

        if tab_run_idx is not None and tab_run_idx > 0:
            # Standard two-part layout: run[0] = left content, tab-run = \t + right
            para.runs[0].text = left
            for k in range(1, tab_run_idx):
                para.runs[k].text = ""
            para.runs[tab_run_idx].text = '\t' + right
            for k in range(tab_run_idx + 1, len(para.runs)):
                para.runs[k].text = ""
            return

    # Default: everything in run[0], silence the rest
    para.runs[0].text = text
    for r in para.runs[1:]:
        r.text = ""

def _clone_para(src_para, text: str):
    new_p = copy.deepcopy(src_para._element)
    for r in new_p.findall(qn("w:r")):
        new_p.remove(r)
    new_r = OxmlElement("w:r")
    orig = src_para._element.findall(qn("w:r"))
    if orig:
        rpr = orig[0].find(qn("w:rPr"))
        if rpr is not None:
            new_r.append(copy.deepcopy(rpr))
    new_t = OxmlElement("w:t")
    new_t.text = text
    if text != text.strip():
        new_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    new_r.append(new_t)
    new_p.append(new_r)
    return new_p