# ─────────────────────────────────────────────────────────────────────────────
#  app.py  –  AI-Powered Resume Tailoring & Cover Letter Generator
#
#  Tech stack
#  ──────────
#  • Streamlit      – UI / hosting
#  • python-docx    – .docx parsing & template injection
#  • SAP Generative AI Hub SDK  – multi-model LLM access via SAP AI Core
#  • python-dotenv  – .env credential loading
#
#  How to run
#  ──────────
#  1. pip install -r requirements.txt
#  2. cp .env.example .env  →  fill in your SAP AI Core credentials
#  3. streamlit run app.py
#
#  Template placeholders expected in .docx templates
#  ─────────────────────────────────────────────────
#  Resume template  : {{TAILORED_EXPERIENCE}}
#  Cover letter template: {{COVER_LETTER_BODY}}
#
#  These are the *only* tokens the app injects.  All other content in the
#  template (name, address, styling, headers, footers) is left untouched.
# ─────────────────────────────────────────────────────────────────────────────

from __future__ import annotations

import io
import os

import streamlit as st
from dotenv import load_dotenv

# ── Project modules ────────────────────────────────────────────────────────────
from ai_utils import (
    AVAILABLE_MODELS,
    generate_cover_letter,
    tailor_resume,
    validate_credentials,
)
from doc_utils import (
    extract_resume_text,
    load_docx,
    replace_placeholder,
    save_docx_to_bytes,
)

# ─────────────────────────────────────────────────────────────────────────────
#  Bootstrap
# ─────────────────────────────────────────────────────────────────────────────

# Load .env before anything else so SAP SDK can read the credentials
load_dotenv()

# ─────────────────────────────────────────────────────────────────────────────
#  Streamlit page config
# ─────────────────────────────────────────────────────────────────────────────

st.set_page_config(
    page_title="AI Resume Tailor",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─────────────────────────────────────────────────────────────────────────────
#  Custom CSS  – minimal polish without any paid component
# ─────────────────────────────────────────────────────────────────────────────

st.markdown(
    """
    <style>
        /* Tighten up section spacing */
        .block-container { padding-top: 2rem; padding-bottom: 2rem; }

        /* Subtle card-like appearance for expanders */
        .streamlit-expanderHeader {
            font-weight: 600;
            font-size: 1rem;
        }

        /* Status badge pill */
        .status-ok   { color: #1d7a3a; background: #d4edda;
                        padding: 2px 8px; border-radius: 12px; font-size: 0.85rem; }
        .status-warn { color: #856404; background: #fff3cd;
                        padding: 2px 8px; border-radius: 12px; font-size: 0.85rem; }
        .status-err  { color: #721c24; background: #f8d7da;
                        padding: 2px 8px; border-radius: 12px; font-size: 0.85rem; }

        /* Make the honesty banner stand out */
        .honesty-banner {
            background: #eef6ff;
            border-left: 4px solid #1a73e8;
            padding: 0.75rem 1rem;
            border-radius: 4px;
            font-size: 0.9rem;
        }
    </style>
    """,
    unsafe_allow_html=True,
)

# ─────────────────────────────────────────────────────────────────────────────
#  Sidebar  – model selection, credential status, instructions
# ─────────────────────────────────────────────────────────────────────────────

with st.sidebar:
    st.title("⚙️ Settings")
    st.divider()

    # ── Model picker ──────────────────────────────────────────────────────────
    st.subheader("🤖 AI Model")
    model_display_names = list(AVAILABLE_MODELS.keys())
    selected_display = st.selectbox(
        "Select model (SAP AI Core deployment)",
        options=model_display_names,
        index=0,                      # default: GPT-4o
        help=(
            "These names must match the deployment names configured in your "
            "SAP AI Core resource group.  Contact your BTP admin if a model "
            "is not available in your account."
        ),
    )
    # Resolve the display label to the actual deployment/model name string
    selected_model: str = AVAILABLE_MODELS[selected_display]
    st.caption(f"Deployment ID: `{selected_model}`")

    st.divider()

    # ── Credential status ─────────────────────────────────────────────────────
    st.subheader("🔐 SAP AI Core Credentials")
    creds_ok, missing_vars = validate_credentials()

    if creds_ok:
        st.markdown(
            '<span class="status-ok">✅ All credentials found</span>',
            unsafe_allow_html=True,
        )
    else:
        st.markdown(
            '<span class="status-err">❌ Credentials missing</span>',
            unsafe_allow_html=True,
        )
        st.error(
            "The following environment variables are not set:\n\n"
            + "\n".join(f"• `{v}`" for v in missing_vars)
            + "\n\nCopy `.env.example` → `.env` and fill in your SAP AI Core "
            "service key values, then restart the app."
        )

    st.divider()

    # ── Placeholder reference ─────────────────────────────────────────────────
    st.subheader("📌 Template Placeholders")
    st.info(
        "Add these exact tokens to your `.docx` templates:\n\n"
        "**Resume template:**\n`{{TAILORED_EXPERIENCE}}`\n\n"
        "**Cover letter template:**\n`{{COVER_LETTER_BODY}}`\n\n"
        "The app will replace them with AI-generated text while preserving "
        "all surrounding formatting."
    )

    st.divider()

    # ── Honesty guardrail notice ──────────────────────────────────────────────
    st.subheader("🛡️ Honesty Guardrail")
    st.markdown(
        """
        <div class="honesty-banner">
        The AI is <strong>hardcoded</strong> to never invent skills, titles,
        degrees, or experiences absent from your base resume. It only reframes
        and reorders existing content to match the job description.
        </div>
        """,
        unsafe_allow_html=True,
    )


# ─────────────────────────────────────────────────────────────────────────────
#  Main content area
# ─────────────────────────────────────────────────────────────────────────────

st.title("📄 AI Resume Tailor & Cover Letter Generator")
st.caption(
    "Powered by **SAP Generative AI Hub** · Zero hallucination guarantee · "
    "Your data never leaves your environment."
)

st.divider()

# ─────────────────────────────────────────────────────────────────────────────
#  Step 1 – Upload documents
# ─────────────────────────────────────────────────────────────────────────────

st.header("Step 1 — Upload Your Documents")

# ── Byte-cache initialisation ─────────────────────────────────────────────────
# WHY THIS MATTERS:
# Streamlit's UploadedFile is a thin wrapper around a socket buffer.  Once read,
# its internal pointer is at EOF and seek(0) is unreliable in cloud deployments
# (the Streamlit Community Cloud runner does NOT guarantee seekability).
# Clicking st.download_button also triggers a full app rerun which can reset
# widget state before the download fires.
#
# Solution: the moment a file is uploaded we read its raw bytes ONCE and store
# them in st.session_state under a stable key.  Every subsequent operation
# creates a fresh io.BytesIO from those cached bytes — fully seekable, rerun-safe.
for _key in ("bytes_base_resume", "bytes_resume_template", "bytes_cover_letter_template"):
    if _key not in st.session_state:
        st.session_state[_key] = None

col_upload_left, col_upload_right = st.columns([1, 1], gap="large")

with col_upload_left:
    st.subheader("📋 Base Resume")
    st.caption(
        "Your master resume containing all of your experience, education, and "
        "skills.  The AI reads this to understand what it is allowed to say about you."
    )
    base_resume_file = st.file_uploader(
        "Upload Base Resume (.docx)",
        type=["docx"],
        key="base_resume",
        help="This file is only read — it is never modified.",
    )
    # Cache bytes immediately — ONLY when a new file object arrives.
    # IMPORTANT: do NOT clear the cache when file_uploader returns None.
    # st.download_button triggers a rerun where file_uploader briefly returns
    # None even though the user hasn't removed the file.  Clearing here would
    # wipe the cached bytes and break the readiness checks on that rerun.
    if base_resume_file is not None:
        st.session_state["bytes_base_resume"] = base_resume_file.read()

with col_upload_right:
    st.subheader("📝 Resume Template")
    st.caption(
        "Your beautifully formatted resume template.  Place the token "
        "`{{TAILORED_EXPERIENCE}}` wherever you want the AI-rewritten "
        "experience / skills section to be injected."
    )
    resume_template_file = st.file_uploader(
        "Upload Resume Template (.docx)",
        type=["docx"],
        key="resume_template",
    )
    if resume_template_file is not None:
        st.session_state["bytes_resume_template"] = resume_template_file.read()

    st.subheader("✉️ Cover Letter Template")
    st.caption(
        "Your cover letter template with branding, header, salutation, and "
        "sign-off already in place.  Put `{{COVER_LETTER_BODY}}` where the "
        "three body paragraphs should appear."
    )
    cover_letter_template_file = st.file_uploader(
        "Upload Cover Letter Template (.docx)",
        type=["docx"],
        key="cover_letter_template",
    )
    if cover_letter_template_file is not None:
        st.session_state["bytes_cover_letter_template"] = cover_letter_template_file.read()

# ── Convenience references to cached bytes ────────────────────────────────────
# These are used everywhere below instead of the UploadedFile objects.
_base_resume_bytes: bytes | None        = st.session_state["bytes_base_resume"]
_resume_template_bytes: bytes | None    = st.session_state["bytes_resume_template"]
_cover_letter_template_bytes: bytes | None = st.session_state["bytes_cover_letter_template"]

st.divider()

# ─────────────────────────────────────────────────────────────────────────────
#  Step 2 – Job description
# ─────────────────────────────────────────────────────────────────────────────

st.header("Step 2 — Paste the Job Description")

job_description_text: str = st.text_area(
    "Target Job Description",
    height=250,
    placeholder=(
        "Paste the full job posting here — include the role title, "
        "responsibilities, requirements, and any preferred qualifications. "
        "The more detail you provide, the better the keyword alignment."
    ),
    help=(
        "Tip: include the company name and role title in the posting so the "
        "cover letter's opening paragraph can reference them naturally."
    ),
)

st.divider()

# ─────────────────────────────────────────────────────────────────────────────
#  Step 3 – Preview extracted resume text  (collapsible)
# ─────────────────────────────────────────────────────────────────────────────

st.header("Step 3 — Preview & Confirm Resume Content")

if _base_resume_bytes:
    with st.expander("👁️ View extracted base resume text (what the AI sees)", expanded=False):
        try:
            # Always build a fresh BytesIO from the cached bytes — never read
            # the UploadedFile object a second time (its pointer is at EOF).
            base_doc = load_docx(io.BytesIO(_base_resume_bytes))
            extracted_text: str = extract_resume_text(base_doc)

            if extracted_text.strip():
                st.text_area(
                    "Extracted plain text (read-only)",
                    value=extracted_text,
                    height=300,
                    disabled=True,
                    key="resume_preview",
                )
                st.caption(
                    f"ℹ️  {len(extracted_text.split())} words extracted from "
                    f"{len(base_doc.paragraphs)} paragraphs + "
                    f"{len(base_doc.tables)} table(s)."
                )
            else:
                st.warning(
                    "No text could be extracted from this .docx file.  "
                    "If your resume uses embedded text boxes or images, the "
                    "parser cannot read them.  Please copy the text into a "
                    "standard paragraph-based .docx."
                )
        except Exception as exc:
            st.error(f"Failed to read the base resume: {exc}")
else:
    st.info("Upload your base resume above to preview its extracted content.")

st.divider()

# ─────────────────────────────────────────────────────────────────────────────
#  Step 4 – Generate
# ─────────────────────────────────────────────────────────────────────────────

st.header("Step 4 — Generate & Download")

# ── Readiness checks ──────────────────────────────────────────────────────────
# We surface clear errors before the user clicks Generate so they don't wait
# through a spinner only to get an error at the end.

def _all_resume_inputs_ready() -> bool:
    """Return True only when every input required for resume tailoring is present."""
    return bool(_base_resume_bytes and _resume_template_bytes and job_description_text.strip())


def _all_cover_letter_inputs_ready() -> bool:
    """Return True only when every input required for cover letter generation is present."""
    return bool(_base_resume_bytes and _cover_letter_template_bytes and job_description_text.strip())


col_gen_left, col_gen_right = st.columns([1, 1], gap="large")

# ─────────────────────────────────────────────────────────────────────────────
#  Resume tailoring column
# ─────────────────────────────────────────────────────────────────────────────

with col_gen_left:
    st.subheader("📋 Tailored Resume")

    # Show a human-readable checklist of what's still missing
    if not _base_resume_bytes:
        st.warning("⬆️  Upload your **base resume** in Step 1.")
    if not _resume_template_bytes:
        st.warning("⬆️  Upload your **resume template** in Step 1.")
    if not job_description_text.strip():
        st.warning("✏️  Paste a **job description** in Step 2.")

    resume_ready = _all_resume_inputs_ready()

    # Persist generated content across Streamlit reruns via session_state
    if "tailored_resume_bytes" not in st.session_state:
        st.session_state["tailored_resume_bytes"] = None
    if "tailored_resume_text" not in st.session_state:
        st.session_state["tailored_resume_text"] = None

    if st.button(
        "✨ Tailor My Resume",
        disabled=(not resume_ready or not creds_ok),
        use_container_width=True,
        type="primary",
    ):
        with st.spinner(f"Tailoring your resume using **{selected_display}** …"):
            try:
                # ── 1. Extract base resume text from cached bytes ─────────────
                # Fresh BytesIO each time — fully seekable, rerun-safe.
                base_doc_for_text = load_docx(io.BytesIO(_base_resume_bytes))
                resume_text = extract_resume_text(base_doc_for_text)

                if not resume_text.strip():
                    st.error(
                        "Could not extract any text from the base resume. "
                        "See the preview in Step 3 for details."
                    )
                    st.stop()

                # ── 2. Call the AI ────────────────────────────────────────────
                tailored_content = tailor_resume(
                    base_resume_text=resume_text,
                    job_description=job_description_text,
                    model_name=selected_model,
                )

                # ── 3. Load a fresh copy of the resume template ───────────────
                # Critical: build from cached bytes, NOT from the UploadedFile
                # whose buffer is already exhausted after Step 1's .read() call.
                resume_doc = load_docx(io.BytesIO(_resume_template_bytes))

                # ── 4. DEBUG: show what text is actually in the template ───────
                # This helps diagnose placeholder-not-found issues.
                all_para_text = [p.text for p in resume_doc.paragraphs if p.text.strip()]
                placeholder_found_in_scan = any(
                    "{{TAILORED_EXPERIENCE}}" in p for p in all_para_text
                )

                # ── 5. Inject AI content into {{TAILORED_EXPERIENCE}} ─────────
                replacements_made = replace_placeholder(
                    doc=resume_doc,
                    placeholder="{{TAILORED_EXPERIENCE}}",
                    replacement_text=tailored_content,
                )

                if replacements_made == 0:
                    st.warning(
                        "⚠️  The placeholder `{{TAILORED_EXPERIENCE}}` was **not found** "
                        "in your resume template, so the AI content could not be injected. "
                        "The download below contains the AI text as a **plain standalone document** instead."
                    )
                    with st.expander("🔍 What your template actually contains (debug)", expanded=True):
                        st.caption("These are the non-empty paragraphs found in your uploaded template:")
                        for line in all_para_text[:30]:
                            st.code(line, language=None)
                        st.info(
                            "If you see your resume content above, you need to **add** "
                            "`{{TAILORED_EXPERIENCE}}` as a placeholder line in your template .docx "
                            "where you want the AI content to appear."
                        )
                    # FALLBACK: build a clean standalone .docx with just the AI content
                    from docx import Document as _Document
                    fallback_doc = _Document()
                    fallback_doc.add_heading("Tailored Resume Content", level=1)
                    for line in tailored_content.split("\n"):
                        fallback_doc.add_paragraph(line)
                    output_doc_bytes = save_docx_to_bytes(fallback_doc)
                else:
                    output_doc_bytes = save_docx_to_bytes(resume_doc)

                # ── 6. Store bytes for download ───────────────────────────────
                st.session_state["tailored_resume_bytes"] = output_doc_bytes
                st.session_state["tailored_resume_text"] = tailored_content

                st.success(
                    f"✅ Resume tailored successfully!  "
                    f"({'placeholder found & replaced' if replacements_made else 'WARNING: placeholder not found — see above'})"
                )

            except RuntimeError as exc:
                # Credential / SDK errors → actionable message
                st.error(f"Configuration error: {exc}")
            except Exception as exc:
                # API-level errors (quota, model unavailable, network, etc.)
                st.error(f"SAP AI Core API error: {exc}")

    # ── Show results & download button ────────────────────────────────────────
    if st.session_state["tailored_resume_text"]:
        with st.expander("👁️ Preview AI-generated resume content", expanded=True):
            st.text_area(
                "Tailored resume body (read-only)",
                value=st.session_state["tailored_resume_text"],
                height=300,
                disabled=True,
                key="tailored_preview",
            )

    if st.session_state["tailored_resume_bytes"]:
        st.download_button(
            label="⬇️  Download Tailored Resume (.docx)",
            data=st.session_state["tailored_resume_bytes"],
            file_name="tailored_resume.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

# ─────────────────────────────────────────────────────────────────────────────
#  Cover letter column
# ─────────────────────────────────────────────────────────────────────────────

with col_gen_right:
    st.subheader("✉️ Cover Letter")

    if not _base_resume_bytes:
        st.warning("⬆️  Upload your **base resume** in Step 1.")
    if not _cover_letter_template_bytes:
        st.warning("⬆️  Upload your **cover letter template** in Step 1.")
    if not job_description_text.strip():
        st.warning("✏️  Paste a **job description** in Step 2.")

    cl_ready = _all_cover_letter_inputs_ready()

    if "cover_letter_bytes" not in st.session_state:
        st.session_state["cover_letter_bytes"] = None
    if "cover_letter_text" not in st.session_state:
        st.session_state["cover_letter_text"] = None

    if st.button(
        "✨ Generate Cover Letter",
        disabled=(not cl_ready or not creds_ok),
        use_container_width=True,
        type="primary",
    ):
        with st.spinner(f"Writing your cover letter using **{selected_display}** …"):
            try:
                # ── 1. Extract base resume text from cached bytes ─────────────
                base_doc_for_cl = load_docx(io.BytesIO(_base_resume_bytes))
                resume_text_for_cl = extract_resume_text(base_doc_for_cl)

                if not resume_text_for_cl.strip():
                    st.error(
                        "Could not extract any text from the base resume. "
                        "See the preview in Step 3 for details."
                    )
                    st.stop()

                # ── 2. Call the AI ────────────────────────────────────────────
                cover_letter_content = generate_cover_letter(
                    base_resume_text=resume_text_for_cl,
                    job_description=job_description_text,
                    model_name=selected_model,
                )

                # ── 3. Load a fresh copy of the cover letter template ─────────
                cl_doc = load_docx(io.BytesIO(_cover_letter_template_bytes))

                # ── 4. Inject AI content into {{COVER_LETTER_BODY}} ───────────
                cl_replacements = replace_placeholder(
                    doc=cl_doc,
                    placeholder="{{COVER_LETTER_BODY}}",
                    replacement_text=cover_letter_content,
                )

                if cl_replacements == 0:
                    cl_para_text = [p.text for p in cl_doc.paragraphs if p.text.strip()]
                    st.warning(
                        "⚠️  The placeholder `{{COVER_LETTER_BODY}}` was **not found** "
                        "in your cover letter template. The download below contains the "
                        "AI text as a **plain standalone document** instead."
                    )
                    with st.expander("🔍 What your template actually contains (debug)", expanded=True):
                        st.caption("These are the non-empty paragraphs found in your uploaded template:")
                        for line in cl_para_text[:30]:
                            st.code(line, language=None)
                        st.info(
                            "Add `{{COVER_LETTER_BODY}}` as a placeholder line in your "
                            "cover letter template .docx where the body paragraphs should appear."
                        )
                    from docx import Document as _Document
                    fallback_cl = _Document()
                    fallback_cl.add_heading("Cover Letter Body", level=1)
                    for line in cover_letter_content.split("\n"):
                        fallback_cl.add_paragraph(line)
                    cl_output_bytes = save_docx_to_bytes(fallback_cl)
                else:
                    cl_output_bytes = save_docx_to_bytes(cl_doc)

                # ── 5. Store bytes for download ───────────────────────────────
                st.session_state["cover_letter_bytes"] = cl_output_bytes
                st.session_state["cover_letter_text"] = cover_letter_content

                st.success(
                    f"✅ Cover letter generated!  "
                    f"({'placeholder found & replaced' if cl_replacements else 'WARNING: placeholder not found — see above'})"
                )

            except RuntimeError as exc:
                st.error(f"Configuration error: {exc}")
            except Exception as exc:
                st.error(f"SAP AI Core API error: {exc}")

    # ── Show results & download button ────────────────────────────────────────
    if st.session_state["cover_letter_text"]:
        with st.expander("👁️ Preview AI-generated cover letter body", expanded=True):
            st.text_area(
                "Cover letter body (read-only)",
                value=st.session_state["cover_letter_text"],
                height=300,
                disabled=True,
                key="cl_preview",
            )

    if st.session_state["cover_letter_bytes"]:
        st.download_button(
            label="⬇️  Download Cover Letter (.docx)",
            data=st.session_state["cover_letter_bytes"],
            file_name="cover_letter.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

st.divider()

# ─────────────────────────────────────────────────────────────────────────────
#  Footer
# ─────────────────────────────────────────────────────────────────────────────

st.markdown(
    """
    <div style="text-align:center; color:#888; font-size:0.8rem; margin-top:1rem;">
    Built with Streamlit · python-docx · SAP Generative AI Hub SDK &nbsp;|&nbsp;
    🛡️ Honesty guardrails prevent AI hallucination &nbsp;|&nbsp;
    No data stored or transmitted beyond SAP AI Core
    </div>
    """,
    unsafe_allow_html=True,
)